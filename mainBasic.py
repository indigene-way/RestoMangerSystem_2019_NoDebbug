#! /usr/bin/python
# -*- coding: utf-8 -*-

import sys, os
import re 
import uuid
import time

from PyQt5 import QtCore, QtGui, uic
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtSql import *

from docx import *

import sqlite3

#===========================================================================DATA BASE 

# base de données sqlite3
conn = sqlite3.connect('dataBase.db')

query = conn.cursor()
				
#===========================================================================TABLE LOGIN
try:
	query.execute("SELECT id FROM Login ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Login")
	conn.execute("""CREATE TABLE Login (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					user_name VARCHAR(50) ,
					user_password VARCHAR(30))""")
					
	# print("DataBase  Login created succefully")
	query.execute("INSERT INTO Login (user_name, user_password) VALUES ('admin','admin')")
				

#===========================================================================TABLE user
try:
	query.execute("SELECT user FROM User ")
except:
	# query.execute("DROP TABLE User")
	conn.execute("""CREATE TABLE User (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					user VARCHAR(50))""")
					
	# print("DataBase  User created succefully")
	query.execute("INSERT INTO User(user) VALUES('')")
								
#===========================================================================TABLE SOCIETY
try:
	query.execute("SELECT id FROM Society ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Society")
	conn.execute("""CREATE TABLE Society (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					society_name VARCHAR(50) ,
					society_number VARCHAR(30))""")
					
	# print("DataBase  Society created succefully")
	query.execute("INSERT INTO Society(society_name, society_number) VALUES ('INDIGENE','0549484715')")
				
#===========================================================================TABLE CATEGORY
try:
	query.execute("SELECT id FROM Categories ORDER BY id DESC")
	# print(str(query.fetchone()[0]).strip("(',')"))
except:
	conn.execute("""CREATE TABLE Categories (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					category_name VARCHAR(50) ,
					category_sub_1 VARCHAR(30) ,
					category_sub_2 VARCHAR(30) ,
					category_sub_3 VARCHAR(30) )""")
					
	# print("DataBase  Categories created succefully")
	query.execute("INSERT INTO Categories (category_name) VALUES\
					('Entrées'),\
					('Plats'),\
					('Sandwiches'),\
					('Pizzas'),\
					('Boissons Fraiches'),\
					('Boissons Chaudes'),\
					('Désserts'),\
					('Jus et Smoothies'),\
					('Autres'),\
					('Supplements');")

#===========================================================================TABLE Products

# query.execute("DROP TABLE Products")
try:
	query.execute("SELECT id FROM Products ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Products (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					product_category_name VARCHAR(50),
					product_subCategory VARCHAR(50),
					product_name VARCHAR(50) ,
					product_price INTEGER ,
					product_qnt INTEGER ,
					product_comptable INTEGER ,
					product_stockable INTEGER ,
					product_total INTEGER,
					product_comptable_cst INTEGER)""")

#============================================================= REFERENCEMENT TICKETS
try:
	query.execute("SELECT id FROM Ref ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Ref (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE, 
					ref INTEGER )""")


	query.execute("INSERT INTO Ref (ref) VALUES (0)")
	query.execute("UPDATE Ref SET ref = 0 WHERE id = 1")
	query.execute("UPDATE Ref SET ref = 0 WHERE id = 2")
			
#===========================================================================REGISTER CATEGORY
# query.execute("DROP TABLE Register")
try:
	query.execute("SELECT id FROM Register ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Register (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					register_date DATE,
					register_sum_init INTEGER,
					register_recette_total INTEGER,
					register_depense_total INTEGER,
					register_ajout_total INTEGER )""")
				
#===========================================================================DEPENSE 
# query.execute("DROP TABLE Register_dep")
try:
	query.execute("SELECT id FROM Register_dep ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Register_dep (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					dep_date DATE,
					dep_type VARCHAR(15),
					dep_description VARCHAR(255),
					dep_value INTEGER)""")
					
	# print("DataBase  Register_dep created succefully")
					
#===========================================================================AJOUT 
# query.execute("DROP TABLE Register_add")
try:
	query.execute("SELECT id FROM Register_add ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Register_add (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					add_date DATE,
					add_depo VARCHAR(50),
					add_description VARCHAR(255),
					add_value INTEGER)""")
					
	# print("DataBase  Register_add created succefully")
				
#============================================================= MAC ADRESS
# query.execute("DROP TABLE MAC")
try:
	query.execute("SELECT id FROM MAC ORDER BY id DESC")
except:
	query.execute("""CREATE TABLE MAC (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE,
					Key VARCHAR(50))""")
	# print("MAC CREATED")


	query.execute("INSERT INTO MAC (Key) VALUES ('')")

conn.commit()
#============================================================================================================DESIGN PATTERN EXCEPT FACTORY
class MessageFactory():

	def raiseAdder(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" ajouté(e) avec succés !")
		msg.exec_()	
		
	def raiseModifier(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" modifié(e) avec succés !")
		msg.exec_()	
		
	def raiseDeleter(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" suprimé(e) avec succés !")
		msg.exec_()	
		
	def raiseCaseExcept(self,case):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("CASE OBLIGATOIRE VIDE !")
		msg.setText("Veuillez saisir "+case)
		msg.exec_()
		
	def raisePrintExcept(self,doc):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("ECHEC IMPRESSION !")
		msg.setText("Veuillez fermer le document word : '"+str(doc)+".docx'  !")
		msg.exec_()	
		
	def raiseCharExcept(self):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("CARACTERE INDESIRABLE !")
		msg.setText("Veuillez ne pas utiliser de caracteres spéciaux ?;:',.$*.... !")
		msg.exec_()
		
	def raiseIndefinedExcept(self,erreur):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("ERREUR !")
		msg.setText(erreur)
		msg.exec_()
		
	def raiseStockAlert(self,stock):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("ALERT STOCK !")
		msg.setText("Vérifié votre stock de "+stock)
		msg.exec_()

class ReferenceSelector():
	def __init__(self):
		self.refSelector()
		self.societySelctor()
	
	def refSelector(self):
		query.execute("SELECT ref FROM Ref WHERE id=1")
		self.ref = str(query.fetchone()).strip("(',')")	
		
	def societySelctor(self):
		query.execute("SELECT society_name FROM Society")
		self.society = query.fetchone() 
		self.society = str(self.society).strip("(',')") 
		query.execute("SELECT society_number FROM Society")
		self.numero = query.fetchone() 
		self.numero = str(self.numero).strip("(',')") 
			
#============================================================================================================CONNECTION 		
	
qtValid= "DESIGN/DIALOGS/connecterDialog.ui"
Ui_ConnectDialog, QtBaseClass = uic.loadUiType(qtValid)

class ConnectDialog(QDialog, Ui_ConnectDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_ConnectDialog.__init__(self)
		self.setupUi(self)
		
		self.disconnect.clicked.connect(self.connSlot)
		
	def connSlot(self):
		if self.sender() == self.disconnect :
			self.destroy()
			
	def login(self,window) :
		query.execute("SELECT user_name FROM Login ORDER BY id")
		self.users=[]
		self.users = query.fetchall()
		query.execute("SELECT user_password FROM Login ORDER BY id")
		self.passwords=[]
		self.passwords=query.fetchall()
		i=0
		j=0
		while i < len(self.users) :
			while j < len(self.passwords) :
				if self.userName.text() == str(self.users[i]).strip("(',')")  and self.userPassword.text() == str(self.passwords[j]).strip("(',')") :
					query.execute("UPDATE User SET user='"+self.userName.text()+"' WHERE id = 1")
					conn.commit()
					window.show()
					self.userName.setText("")
					self.userPassword.setText("")
					self.destroy()
				j+=1
				i+=1	
									
#============================================================================================================SETTINGS
qtValid= "DESIGN/DIALOGS/settingCreatorDialog.ui"
Ui_SettingCreatorDialog, QtBaseClass = uic.loadUiType(qtValid)

class SettingCreatorDialog(QDialog, Ui_SettingCreatorDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_SettingCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.referenceSelector = ReferenceSelector()
		
		self.messageFactory = MessageFactory()
		
		self.settingSignal()
		self.selection()
		
	def selection(self):
		self.society = self.referenceSelector.society 
		self.numero = self.referenceSelector.numero  
		
		self.societyName.setText(self.society)
		self.societyNumber.setText(self.numero)
		
		query.execute("SELECT user FROM User WHERE id = 1")
		self.user = query.fetchone()
		self.user= str(self.user).strip("(',')")
		self.who.setText(self.user)
		
		query.execute("SELECT user_name FROM Login ORDER BY id ASC")
		self.users=[]
		self.users = query.fetchall()
		query.execute("SELECT user_password FROM Login ORDER BY id ASC")
		self.passwords=[]
		self.passwords=query.fetchall()
		
		self.modSetComboBox.clear()
		self.modSetComboBox.addItem("...")
		self.delSetComboBox.clear()
		self.delSetComboBox.addItem("...")
		
		i=0
		for user in self.users :
			i+=1
			self.userNbr.setText(str(i))
			self.modSetComboBox.addItem(str(user).strip("(',')"))
			self.delSetComboBox.addItem(str(user).strip("(',')"))
					
	def settingSignal(self) :
		self.logAdd.clicked.connect(self.settingSlot)
		
		self.logModif.clicked.connect(self.settingSlot)
		self.modSetComboBox.currentTextChanged.connect(self.comboSetting)
		
		self.logDel.clicked.connect(self.settingSlot)
		self.delSetComboBox.currentTextChanged.connect(self.comboSetting)
		
		self.settingInfo.clicked.connect(self.settingSlot)
		
		self.setOK.clicked.connect(self.settingSlot)
		self.setAnnuler.clicked.connect(self.settingSlot)
		
	def settingSlot(self):
		if self.sender() == self.logAdd :
			self.settingAdd()
			self.selection()
			
		if self.sender() == self.logModif :
			self.settingModif()
			self.selection()
			
		if self.sender() == self.logDel :
			self.settingDel()
			self.selection()			
		
		if self.sender() == self.settingInfo :
			self.socityInfo()
			self.selection()		
		
		if self.sender() == self.setOK :
			self.selection()	
			self.destroy()
		
		if self.sender() == self.setAnnuler :	
			self.destroy()
			
	def comboSetting(self):
			
		self.modUserName.setText(self.modSetComboBox.currentText())
		self.delUserName.setText(self.delSetComboBox.currentText())
		
	def settingAdd(self) :
		try:
			if self.newUserName.text() != "" and self.newPassWord.text() != "" :
				query.execute("INSERT INTO Login (user_name, user_password) VALUES('"+self.newUserName.text()+"','"+self.newPassWord.text()+"')")
				conn.commit()
				self.newUserName.setText("")
				self.newPassWord.setText("")
				
				self.messageFactory.raiseAdder("Utilisateur")
				self.selection()
				
			else :
				self.messageFactory.raiseCaseExcept("toute les case")
				if self.newUserName.text() != "" :
					self.messageFactory.raiseCaseExcept("le nom d'utilisaeur")
				if self.newPassWord.text() != "":
					self.messageFactory.raiseCaseExcept("le mot de passe")
		except:
			self.messageFactory.raiseCharExcept()
		
	def settingModif(self) :
		try:
			if self.modUserName.text() != "" and self.modPassWord.text() != "" and self.who.text() == self.modSetComboBox.currentText() and \
				self.modUserName.text() != "..." and self.modSetComboBox.currentText() != "..." :
				query.execute("UPDATE Login SET user_name='"+self.modUserName.text()+"', user_password='"+self.modPassWord.text()+"' WHERE user_name = '"+self.modSetComboBox.currentText()+"'")
				conn.commit()
				
				self.messageFactory.raiseModifier("Utilisateur")
				
				self.modUserName.setText("")
				self.modPassWord.setText("")
				self.selection()
				
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.modUserName.text() == "" :
					self.messageFactory.raiseCaseExcept("le nom d'utilisateur")
				if self.modPassWord.text() == "" :
					self.messageFactory.raiseCaseExcept("le mot de passe")
				# elif self.modSetComboBox.currentText() != self.user :
					# msg.setText("seul "+self.modSetComboBox.currentText()+" est abilité à modifier se compte !")
					
				self.modUserName.setText("")
				self.modPassWord.setText("")
		except:
			self.messageFactory.raiseCharExcept()
			
	def settingDel(self) :
		try:
			if self.delUserName.text() != "" and self.delUserName.text() != "..." and self.delSetComboBox.currentText() != "..."  and\
			self.delSetComboBox.currentText() == self.user :
				query.execute("DELETE FROM 'Login' WHERE user_name = '"+self.delSetComboBox.currentText()+"'")
				conn.commit()
				
				self.messageFactory.raiseDeleter("Utilisateur")
				
				self.delUserName.setText("")
				self.selection()
				
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.delUserName.text() == "" :
					self.messageFactory.raiseCaseExcept("le Nom d'utilisateur")
				# if self.user != "admin" :
					# msg.setText("seul admin est abilité a changer les comptes !")
				# elif self.delSetComboBox.currentText() != self.user :
					# msg.setText("seul "+self.delSetComboBox.currentText()+" est abilité à modifier se compte !")
					
				self.delUserName.setText("")
		except:
			self.messageFactory.raiseCharExcept()
		
	def socityInfo(self):
		try:
			if self.societyName.text() != "" and self.societyNumber.text() != "" :
				query.execute("UPDATE Society SET society_name='"+self.societyName.text()+"', society_number='"+self.societyNumber.text()+"' WHERE id = 1 ")
				conn.commit()
							
				self.messageFactory.raiseModifier("Information")
				
				self.modUserName.setText("")
				self.modPassWord.setText("")
				
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.societyName.text() == "" :
					self.messageFactory.raiseCaseExcept("le nom de la society")
				elif self.societyNumber.text() == "" :
					self.messageFactory.raiseCaseExcept("le numéro de la society")
		except:
			self.messageFactory.raiseCharExcept()
	

#================================================================================= PRODUCT WIDGET
#=============VLD DIALOG		
qtValid= "DESIGN/DIALOGS/VLDDialog.ui"
Ui_validDialog, QtBaseClass = uic.loadUiType(qtValid)

class ValidProdsDialog(QDialog, Ui_validDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_validDialog.__init__(self)
		self.setupUi(self)


#=============CATEGORY DIALOG
qtCreatorDialog= "DESIGN/DIALOGS/categoryCreatorDialog.ui"
Ui_categoryCreatorDialog, QtBaseClass = uic.loadUiType(qtCreatorDialog)

class CategoryCreatorDialog(QDialog, Ui_categoryCreatorDialog):# EDIT : MODIF Product Catégory DIALOG

	def __init__(self):
		QDialog.__init__(self)
		Ui_categoryCreatorDialog.__init__(self)
		self.setupUi(self)
		self.messageFactory = MessageFactory()
		# self.prodWidget = ProductWidget()	
		self.setWindowTitle("Edition du menu et sous-catégories")
		# SETTING TEXT  category_name FROM CATEGORIES -->  CATLIST
		
		#GETTING CATEGORY NAMES FROM CATEGORY QDIALOG INTO LISTS
		#CATEGORY PRODUCTS NAMES LIST
		self.categoryName = [self.categoryName1, self.categoryName2, self.categoryName3,
		self.categoryName4,  self.categoryName5,  self.categoryName6, self.categoryName7,
		self.categoryName8,  self.categoryName9, self.categoryName10]
  
		#CATEGORY PRODUCTS SUB NAMES 1 LIST
		self.categorySub1 = [self.cat1subCat1, self.cat2subCat1, self.cat3subCat1,
		self.cat4subCat1, self.cat5subCat1, self.cat6subCat1, self.cat7subCat1, 
		self.cat8subCat1, self.cat9subCat1, self.cat10subCat1]

		#CATEGORY PRODUCTS SUB NAMES 2 LIST
		self.categorySub2 = [self.cat1subCat2, self.cat2subCat2, self.cat3subCat2,
		self.cat4subCat2, self.cat5subCat2, self.cat6subCat2, self.cat7subCat2, 
		self.cat8subCat2, self.cat9subCat2, self.cat10subCat2 ]


		#CATEGORY PRODUCTS SUB NAMES 3 LIST
		self.categorySub3 = [self.cat1subCat3, self.cat2subCat3, self.cat3subCat3,
		self.cat4subCat3, self.cat5subCat3, self.cat6subCat3, self.cat7subCat3, 
		self.cat8subCat3, self.cat9subCat3, self.cat10subCat3]	
		
		# query.execute("SELECT * FROM products")
		# print(query.fetchall())
		
		# query.execute("SELECT id, category_name FROM categories")
		# print(query.fetchall())
		
			
		# SELECTING CATEGORIES FROM DATA --> INTO LIST 
		
		# SELECTING  category_name FROM CATEGORIES -->  CATLIST:
		query.execute("SELECT category_name FROM Categories ORDER by id ASC ")#SELECT : Category Name
		self.catList = list()
		self.catList = query.fetchall()
		i=0
		while i <= 9 :
			self.categoryName[i].setText(str(self.catList[i]).strip("(',')"))
			self.categorytabWidget.setTabText(i,str(self.catList[i]).strip("(',')"))#TAB 1 SETTEXT
			i+=1
	
		
		query.execute("SELECT category_sub_1 FROM Categories ORDER by id ASC ")#SELECT : Category sub 1
		self.subList1 = list()
		self.subList1 = query.fetchall()
		i=0
		while i <= 9 :
			if str(self.subList1[i]).strip("(',')") != "None" :
				self.categorySub1[i].setText(str(self.subList1[i]).strip("(',')"))
			i+=1

			
		query.execute("SELECT category_sub_2 FROM Categories ORDER by id ASC ")#SELECT : Category sub 2
		self.subList2 = list()
		self.subList2 = query.fetchall()
		i=0
		while i <= 9 :
			if str(self.subList2[i]).strip("(',')") != "None" :
				self.categorySub2[i].setText(str(self.subList2[i]).strip("(',')"))
			i+=1

			
		query.execute("SELECT category_sub_3 FROM Categories ORDER by id ASC ")#SELECT : Category sub 3
		self.subList3 = list()
		self.subList3 = query.fetchall()
		i=0
		while i <= 9 :
			if str(self.subList3[i]).strip("(',')") != "None" :
				self.categorySub3[i].setText(str(self.subList3[i]).strip("(',')"))
			i+=1
			
		#modif products cat_name

	def AnnulerNewCategorySlot(self):#SELF.DIALOG.DESTROY
		self.destroy()
		
	def validNewCategorySlot(self):#UPDATE * SET Categories Name and Subs BY id, UPDATE SELF.CATNAMES, TABSNAMES in DIALOG AND MAIN
		
		#SELECTING PRODUCT CATEGORY NAMES FROM PRODUCT AND UPDATE FROM CATLIST
		# query.execute("SELECT product_category_name FROM Products ORDER BY id ASC")
		query.execute("SELECT product_category_name FROM Products ORDER BY id")
		prodcatlist = query.fetchall()
		query.execute("SELECT product_subCategory FROM Products ORDER BY id")
		prodsublist = query.fetchall()
		# for i in suberlist:
			# print(i)	
		
		try:
			c=0
			while c <= 9 :
			
				#CATEGORY  UPDATING
				query.execute("UPDATE Categories SET category_name = '"+self.categoryName[c].text()+"' WHERE id = "+str(c+1)+"")
						
				query.execute("UPDATE Categories SET category_sub_1 = '"+self.categorySub1[c].text()+"' WHERE id = "+str(c+1)+"")
					
				query.execute("UPDATE Categories SET category_sub_2 = '"+self.categorySub2[c].text()+"' WHERE id = "+str(c+1)+"")

				query.execute("UPDATE Categories SET category_sub_3 = '"+self.categorySub3[c].text()+"' WHERE id = "+str(c+1)+"") 

			
				#UPDATING PRODUCTS CATEGORY NAMES WITH NEW DATA
				
				for prod in prodcatlist:
					if str(prod).strip("(',')") == str(self.catList[c]).strip("(',')") :
						query.execute("UPDATE Products SET product_category_name = '"+self.categoryName[c].text()+"' WHERE product_category_name ='"+str(self.catList[c]).strip("(',')")+"'")
				
				for prod in prodsublist:
					if str(prod).strip("(',')") == str(self.subList1[c]).strip("(',')") :
						query.execute("UPDATE Products SET product_subCategory = '"+self.categorySub1[c].text()+"' WHERE product_subCategory ='"+str(self.subList1[c]).strip("(',')")+"'")
						query.execute("SELECT * FROM Products WHERE product_subCategory = '"+self.categorySub1[c].text()+"'")
						
					if str(prod).strip("(',')") == str(self.subList2[c]).strip("(',')") :
						query.execute("UPDATE Products SET product_subCategory = '"+self.categorySub2[c].text()+"' WHERE product_subCategory ='"+str(self.subList2[c]).strip("(',')")+"'")
						query.execute("SELECT * FROM Products WHERE product_subCategory = '"+self.categorySub2[c].text()+"'")
						
					if str(prod).strip("(',')") == str(self.subList3[c]).strip("(',')") :
						query.execute("UPDATE Products SET product_subCategory = '"+self.categorySub3[c].text()+"' WHERE product_subCategory ='"+str(self.subList3[c]).strip("(',')")+"'")
						query.execute("SELECT * FROM Products WHERE product_subCategory = '"+self.categorySub3[c].text()+"'")
				c+=1
			# SELECTING CATEGORIES FROM DATA --> INTO LIST 
			
			# SELECTING  category_name FROM CATEGORIES -->  CATLIST:
			query.execute("SELECT category_name FROM Categories ORDER by id ASC ")#SELECT : Category Name
			self.catList = list()
			self.catList = query.fetchall()
			i=0
			while i <= 9 :
				self.categoryName[i].setText(str(self.catList[i]).strip("(',')"))
				self.categorytabWidget.setTabText(i,str(self.catList[i]).strip("(',')"))#TAB 1 SETTEXT
				i+=1

			
			query.execute("SELECT category_sub_1 FROM Categories ORDER by id ASC ")#SELECT : Category sub 1
			self.subList1 = list()
			self.subList1 = query.fetchall()
			i=0
			while i <= 9 :
				if str(self.subList1[i]).strip("(',')") != "None" :
					self.categorySub1[i].setText(str(self.subList1[i]).strip("(',')"))
				i+=1

				
			query.execute("SELECT category_sub_2 FROM Categories ORDER by id ASC ")#SELECT : Category sub 2
			self.subList2 = list()
			self.subList2 = query.fetchall()
			i=0
			while i <= 9 :
				if str(self.subList2[i]).strip("(',')") != "None" :
					self.categorySub2[i].setText(str(self.subList2[i]).strip("(',')"))
				i+=1

				
			query.execute("SELECT category_sub_3 FROM Categories ORDER by id ASC ")#SELECT : Category sub 3
			self.subList3 = list()
			self.subList3 = query.fetchall()
			i=0
			while i <= 9 :
				if str(self.subList3[i]).strip("(',')") != "None" :
					self.categorySub3[i].setText(str(self.subList3[i]).strip("(',')"))	
				i+=1

			conn.commit()
		except:
			self.messageFactory.raiseCharExcept()

#=============PRODUCT CREATOR DIALOG

qtProductCreator= "DESIGN\DIALOGS\productCreatorDialog.ui"
Ui_productCreatorDialog, QtBaseClass = uic.loadUiType(qtProductCreator)

class ProductCreatorDialog(QDialog, Ui_productCreatorDialog):#EDIT : MODIF Product Name,Price DIALOG
	
	def __init__(self):
		QDialog.__init__(self)
		Ui_productCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("Edition des Produits")
		
		self.prodAnnuler.clicked.connect(self.prodEditAnnuler)
		
		#======INSTANCES
		self.categoryDialog = CategoryCreatorDialog()#EDIT : CATEGORY CHANGES DIALOG 
		self.prodDataList()
		#======SLOT VALiD
		
		self.categoryAddComboBox.clear()
		self.categoryModifComboBox.clear()
		self.categoryDelComboBox.clear()
		
		self.categoryAddComboBox.addItem("")
		self.categoryModifComboBox.addItem("")
		self.categoryDelComboBox.addItem("")
		
		self.prodOK.clicked.connect(self.prodEditValid)

		# """ RECUPERATION DE LA CLASSE VIA LISTES"""
		query.execute("SELECT category_name FROM Categories ORDER BY id")
		self.categoryList = list()
		self.categoryList = query.fetchall()

	def prodDataList(self):# i=0
		
		query.execute("SELECT product_name FROM Products ORDER BY id ASC")
		self.prodNameList = list()
		self.prodNameList = query.fetchall()
		
		query.execute("SELECT product_price FROM Products ORDER BY id ASC")
		self.prodPriceList = list()
		self.prodPriceList = query.fetchall()
		
		query.execute("SELECT product_qnt FROM Products ORDER BY id ASC")
		self.prodQntList = list()
		self.prodQntList = query.fetchall()
		
		query.execute("SELECT product_comptable FROM Products ORDER BY id ASC")
		self.prodComptList = list()
		self.prodComptList = query.fetchall()
		
		query.execute("SELECT product_stockable FROM Products ORDER BY id ASC")
		self.prodStockList = list()
		self.prodStockList = query.fetchall()
		
		query.execute("SELECT product_total FROM Products ORDER BY id ASC")
		self.prodTotalList = list()
		self.prodTotalList = query.fetchall()
		
	def prodEditValid(self):
		            		
		self.close()

	def prodEditAnnuler(self): 
		self.close()	
	
#=============INITSOM DIALOG		
qtInitial= "DESIGN/DIALOGS/iniSum.ui"
Ui_initDialog, QtBaseClass = uic.loadUiType(qtInitial)

class InitDialog(QDialog, Ui_initDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_initDialog.__init__(self)
		self.setupUi(self)
		
		#===========LOGICALS
		self.Date = time.strftime("%Y-%m-%d")
		
		self.dataInit()
		
	def dataInit(self) :
		query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
		self.sominit = query.fetchone()
		self.sominit = str(self.sominit).strip("(',')")
		self.sumInit.setText(self.sominit)
		
		self.INIT_OK.clicked.connect(self.slot)
		self.INIT_NO.clicked.connect(self.slot)
		
	def slot(self) :
	
		if self.sender() == self.INIT_OK :
			query.execute("UPDATE Register SET register_sum_init = '"+self.sumInit.text()+"' WHERE register_date = '"+self.Date+"'")
			query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
			self.sominit = query.fetchone()
			self.sominit = str(self.sominit).strip("(',')")
			self.sumInit.setText(self.sominit)
			
			conn.commit()
				 
			msg = QMessageBox()
			msg.setIcon(QMessageBox.Information)

			msg.setText("Enregistrée avec succés")
			msg.setWindowTitle("Edition de la caisse !")
			msg.exec_()
			self.destroy()
		
		if self.sender() == self.INIT_NO :
			self.destroy()

#=============REGISTER DIALOG
qtRegisterDialog= "DESIGN/DIALOGS/registerCreatorDialog.ui"
Ui_registerCreatorDialog, QtBaseClass = uic.loadUiType(qtRegisterDialog)
class RegisterCreatorDialog(QDialog, Ui_registerCreatorDialog):# EDIT : MODIF Product Catégory DIALOG

	def __init__(self):
		QDialog.__init__(self)
		Ui_registerCreatorDialog.__init__(self)
		self.setupUi(self)

		self.setWindowTitle("Gestion de la Caisse et Des Sommes d'Argent")
		self.referenceSelector = ReferenceSelector()
		self.referenceSelector.refSelector()
		self.referenceSelector.societySelctor()
		
		#===========LOGICALS
		self.Date = time.strftime("%Y-%m-%d")
			
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		
		#==========INIT INSTANCES
		self.productDialog = ProductCreatorDialog()
		self.registerInit = InitDialog()
		self.messageFactory = MessageFactory()
		
		self.registerData()
		self.registerSignal()
		
		#===========METHODES
	def registerData(self) :
		
		query.execute("SELECT register_date FROM Register ORDER BY id DESC")
		self.lastDate = query.fetchall()
		
		if self.lastDate == [] :
			query.execute("INSERT INTO Register (register_date, register_sum_init, register_recette_total,register_depense_total,register_ajout_total) VALUES ('"\
			+self.Date+"',0,0,0,0)")
			self.recette = 0
			
		elif str(self.lastDate[0]).strip("(',')") != self.Date :
			query.execute("INSERT INTO Register (register_date, register_sum_init, register_recette_total,register_depense_total,register_ajout_total) VALUES ('"\
			+self.Date+"',0,0,0,0)")
			self.recette = 0
			
		query.execute("SELECT register_recette_total FROM Register WHERE register_date = '"+self.Date+"'")
		self.recette = str(query.fetchone()).strip("(',')")
		self.recette = int(self.recette)
		
		# query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
		query.execute("SELECT register_sum_init FROM Register ORDER BY id DESC")

		query.execute("SELECT register_sum_init FROM Register ORDER BY id DESC")
		self.suminit = str(query.fetchone()).strip("(',')")
		
		query.execute("SELECT register_depense_total FROM Register WHERE register_date = '"+self.Date+"'")
		self.depense = str(query.fetchone()).strip("(',')")
		
		query.execute("SELECT register_ajout_total FROM Register WHERE register_date = '"+self.Date+"'")
		self.ajout = str(query.fetchone()).strip("(',')")
		
		self.sumInit.setText(self.suminit)
		self.sumDay.setText(str(self.recette))
		self.sumDep.setText(self.depense)
		self.sumAdd.setText(str(self.ajout))
		
		self.sumTotal.display(0)
		conn.commit()

	def registerSignal(self):
	
		self.totalButton.clicked.connect(self.registerSlot)
		self.depAdd.clicked.connect(self.registerSlot)
		self.addAdd.clicked.connect(self.registerSlot)
		
		self.registerAnnuler.clicked.connect(self.registerSlot)
		self.registerPrint.clicked.connect(self.registerSlot)
		self.editInit.clicked.connect(self.registerSlot)
		
		self.registerInit.INIT_OK.clicked.connect(self.totalRegister)

	def totalRegister(self):
		try:
			self.regTotal = []
			query.execute("SELECT register_recette_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.recette = str(query.fetchone()).strip("(',')")	
			self.regTotal.append(int(self.recette))
			
			query.execute("SELECT register_sum_init FROM Register ORDER BY id DESC")
			self.suminit = str(query.fetchone()).strip("(',')")
			self.regTotal.append(int(self.suminit))
			
			query.execute("SELECT register_depense_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.depense = str(query.fetchone()).strip("(',')")
			
			query.execute("SELECT register_ajout_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.ajout = str(query.fetchone()).strip("(',')")
			self.regTotal.append(int(self.ajout))
			
			self.sumInit.setText(self.suminit)
			self.sumDay.setText(str(self.recette))
			self.sumDep.setText(self.depense)
			self.sumAdd.setText(str(self.ajout))
			
			self.sumTotal.display(sum(self.regTotal) - int(self.depense))
			self.actual = str(sum(self.regTotal) - int(self.depense))
			self.actual = str(self.actual)	
		except:
			self.messageFactory.raiseIndefinedExcept("au niveau de la caisse")
		
	def registerSlot(self) :	
		self.regTotal = []
		
		if self.sender() == self.totalButton :
			self.totalRegister()
		
		if self.sender() == self.registerAnnuler :
			self.destroy()
		
		if self.sender() == self.depAdd :
			self.newDepense()
			
		if self.sender() == self.editInit :
			self.totalRegister()
			self.registerInit.show()
			
		if self.sender() == self.addAdd :
			self.newDepot()
		
		if self.sender() == self.registerPrint :
			self.totalRegister()
			self.printRegister()

	def newDepot(self):
		try:
			if self.addVal.text() != '0' and self.addName.text() != "" :
				query.execute("INSERT INTO Register_add (add_date, add_depo, add_description,add_value) VALUES ('"\
				+self.Date+"','"+str(self.addName.text())+"','"+str(self.addDesc.text())+"','"+str(self.addVal.text())+"')")
				conn.commit()
				 
				msg = QMessageBox()
				msg.setIcon(QMessageBox.Information)

				msg.setText("Ajout Enregistrée avec succés")
				msg.setWindowTitle("Edition des Ajout !")
				msg.exec_()
		
				self.addName.setText("")
				self.addDesc.setText("")
				self.addVal.setText("0")
		
				query.execute("SELECT add_value FROM Register_add WHERE add_date = '"+self.Date+"'")
				self.addList = []
				self.addList = query.fetchall()
				self.add = []
				i=0
				for add in self.addList :	
					self.add.append(str(add).strip("(',')"))
					self.add[i] = int(self.add[i])
					i+=1
				query.execute("UPDATE Register SET register_ajout_total = "+str(sum(self.add))+" WHERE register_date ='"+self.Date+"'")
				
				query.execute("SELECT * FROM Register")
				self.regTotal.append(sum(self.add))
				conn.commit()
				self.totalRegister()
				
			else : 
				self.messageFactory.raiseCaseExcept("Toutes les cases")
				if  self.addName.text() == '':
					self.messageFactory.raiseCaseExcept("la case Dépositaire")
				if  self.addVal.text() == '0':
					self.messageFactory.raiseCaseExcept("la case Somme à dépenser")
		except:
			self.messageFactory.raiseCharExcept()

	def actualDepense(self):	
		query.execute("SELECT dep_value FROM Register_dep WHERE dep_date = '"+self.Date+"'")
		self.depensesList = []
		self.depensesList = query.fetchall()
		self.dep = []
		i=0
		for dep in self.depensesList :	
			self.dep.append(str(dep).strip("(',')"))
			self.dep[i] = int(self.dep[i])
			i+=1
		query.execute("UPDATE Register SET register_depense_total = "+str(sum(self.dep))+" WHERE register_date ='"+self.Date+"'")
		
		query.execute("SELECT register_depense_total FROM Register WHERE register_date = '"+self.Date+"'")
		self.depense = str(query.fetchone()).strip("(',')")
		self.depense = int(self.depense)
		conn.commit()
		self.totalRegister()
			
	def newDepense(self):
		try:
			if self.depVal.text() != '0' and self.depDesc.text() != "" :
				query.execute("INSERT INTO Register_dep (dep_date, dep_type, dep_description,dep_value) VALUES ('"\
				+self.Date+"','"+self.depTypeCmbBox.currentText()+"','"+self.depDesc.text()+"','"+self.depVal.text()+"')")
				query.execute("SELECT * FROM Register_dep")
				conn.commit()
				 
				
				self.messageFactory.raiseAdder("Dépense")
		
				self.depDesc.setText("")
				self.depVal.setText("0")
				self.totalRegister()
				
				self.actualDepense()
				
			else : 
				self.messageFactory.raiseCaseExcept("Toutes les cases")
				if  self.depDesc.text() == '':
					self.messageFactory.raiseCaseExcept("la case déscription")
				if  self.depVal.text() == '0':
					self.messageFactory.raiseCaseExcept("la case Somme à déposer")
		except:
			self.messageFactory.raiseCharExcept()

	def	printRegister(self) :
		try :			
			TICKET = Document()
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKET.add_heading("                        "+self.referenceSelector.society+" : 'Mouvement de Caisse'", level=1)
			h.bold = True
			h.italic = True
			p = TICKET.add_paragraph("				             "+self.Date)
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("                                     Somme Actuelle en Caisse : "+str(self.actual)+" DA ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#DEPENSES
			p = TICKET.add_heading("DEPENSES : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			query.execute("SELECT dep_type FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT dep_description FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT dep_value FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.valList = []
			self.valList = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Type'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				pi+=1
				
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#AJOUTS		
			p = TICKET.add_heading("DEPOTS : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			query.execute("SELECT add_depo FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT add_description FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT add_value FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.valList = []
			self.valList = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Depositaire'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				pi+=1
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph("SOMME INITIALE EN CAISSE : " +self.suminit+ " DA")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph("RECETTE DU JOUR  : " +self.recette+ " DA")
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("\t\tReference No : "+str(self.referenceSelector.ref)+". --"+self.referenceSelector.society+"-- "+self.localDateTime, level=3)
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			TICKET.save('DOCUMENTS/CAISSE/Caisse'+str(self.referenceSelector.ref)+"_"+self.localDate+'.docx' )
			os.startfile('DOCUMENTS\CAISSE\Caisse'+str(self.referenceSelector.ref)+"_"+self.localDate+'.docx' , 'print')
			
			conn.commit()
			
		except :
			self.messageFactory.raisePrintExcept("Caisse"+str(self.ref)+"_"+self.localDate)

#========== PRODUCT WIDGET UIC CONVERT & Load print

qtProductWidget = "DESIGN/WIDGETS/ProductWidget.ui" # Enter file here.

Ui_ProductWidget, QtBaseClass = uic.loadUiType(qtProductWidget)

#==========INIT CLASS

class ProductWidget(QWidget, Ui_ProductWidget):

	def __init__(self):
	
		QWidget.__init__(self)#QWidget init
		Ui_ProductWidget.__init__(self)#ui init
		self.setupUi(self)#UI init
				
		#==========Logical & Réusables Variables
		self.referenceSelector = ReferenceSelector()
			
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		self.Date = time.strftime("%Y-%m-%d")
		
		self.society = self.referenceSelector.society 
		self.numero = self.referenceSelector.numero  
		
		query.execute("SELECT user FROM User WHERE id = 1")
		self.user = query.fetchone()
		self.user= str(self.user).strip("(',')")
				
		self.ref = self.referenceSelector.ref
		self.viewRefNum.setText(str(self.ref))

		#==========INIT CLASSES OF TYPE DIALOGS FROM MODULE "CreatorDialog":
		
		self.categoryDialog = CategoryCreatorDialog()#EDIT : CATEGORY CHANGES DIALOG
		
		self.productDialog = ProductCreatorDialog()#EDIT : CHANGE Prod Name,Price DIALOG
		
		self.validVLD = ValidProdsDialog()#CONFIRM : TOTAL AND VALID DIALOG TablesNumDialog
		
		self.register = RegisterCreatorDialog()
		
		self.register = RegisterCreatorDialog()
		
		self.messageFactory = MessageFactory()
		
		#==========DELEGATES INSTANCES :
		
		
		self.categoryDataManager()#DATAMANAGER : Categories.db
		
		self.categorySignalClicked()#SIGNAL : EMIT FOR CATEGORY WHEN CLICKED
		
		self.categorySlotClicked()#SLOT : Work ON Category WHEN CLICKED
		
		# self.productsDataManager()#DATAMANAGER : Products.db
		
		self.productsSignalClicked()#SIGNAL : EMIT wheN PRODUCT IS CLICKED
		
		self.validationButtonsSignalClicked()#SIGNAL : EMIT WhEN VALIDATIONBUTTON IS CLICKED
		
		self.firstView()#VIEW : Index View on GUI's openning
#=========================================		
	
	def categoryDataManager(self):	#SELECT Categories.names|subs --> catList|subListi --> menuprod_catName, tabs, dialog(catName, subicatsub)"""
									#UPDATE prod_Category FROM catList by id """
		
		self.menuProd_categoryList=[self.menuProd_category_1, self.menuProd_category_2, self.menuProd_category_3,
				self.menuProd_category_4, self.menuProd_category_5, self.menuProd_category_6, self.menuProd_category_7,
				self.menuProd_category_8, self.menuProd_category_9, self.menuProd_category_10]
#CREATE A prod_List[prod_buttons]
		self.prod_List = list()
		
		self.prod_List = [self.prod_1, self.prod_2, self.prod_3, self.prod_4, self.prod_5, self.prod_6, 
		self.prod_7, self.prod_8, self.prod_9, self.prod_10, self.prod_11, self.prod_12, self.prod_13, 
		self.prod_14, self.prod_15, self.prod_16, self.prod_17, self.prod_18, self.prod_19, self.prod_20, 
		self.prod_21, self.prod_22, self.prod_23, self.prod_24, self.prod_25, self.prod_26, self.prod_27, 
		self.prod_28, self.prod_29, self.prod_30, self.prod_31, self.prod_32, self.prod_33, self.prod_34, 
		self.prod_35, self.prod_36, self.prod_37, self.prod_38, self.prod_39, self.prod_40, self.prod_41, 
		self.prod_42, self.prod_43, self.prod_44, self.prod_45, self.prod_46, self.prod_47, self.prod_48, 
		self.prod_49, self.prod_50, self.prod_51, self.prod_52, self.prod_53, self.prod_54, self.prod_55, 
		self.prod_56, self.prod_57, self.prod_58, self.prod_59, self.prod_60, self.prod_61, self.prod_62, 
		self.prod_63, self.prod_64, self.prod_65, self.prod_66, self.prod_67, self.prod_68, self.prod_69, 
		self.prod_70, self.prod_71, self.prod_72, self.prod_73, self.prod_74, self.prod_75]		
		c=0
		while c <= 9:
		
			self.categoryDialog.catList[c] = self.categoryDialog.categoryName[c].text()
			self.menuProd_categoryList[c].setText(str(self.categoryDialog.catList[c]).strip("(',')"))
			self.menuProd_categoryList[c].setToolTip(str(self.categoryDialog.catList[c]).strip("(',')"))
			self.menuProd_categoryList[c].setStatusTip(str(self.categoryDialog.catList[c]).strip("(',')")+"...")
			c+=1	
			
	def categorieCreatorSignal(self):#SIGNAL : CALL DIALOG ui
		self.categoryDialog.catOK.clicked.connect(self.categoryDialog.validNewCategorySlot)#SIGNAL : VALID CHANGES ON CATEGORIES and SUbCategories	
		self.categoryDialog.catOK.clicked.connect(self.validNewCategorySlot)#SIGNAL : VALID CHANGES ON CATEGORIES and SUbCategories
		self.categoryDialog.catAnnuler.clicked.connect(self.categoryDialog.AnnulerNewCategorySlot)#Valid and Save changes Function
		self.categoryDialog.show()
		
	def validNewCategorySlot(self):#UPDATE * SET Categories Name and Subs BY id, UPDATE SELF.CATNAMES, TABSNAMES in DIALOG AND MAIN
						
		c=0
		while c <= 9 :
		
			#CATEGORY  UPDATING
			
			self.menuProd_categoryList[c].setText(self.categoryDialog.categoryName[c].text())#Menu Category 1 SETTEXT
			self.categoryDialog.catList[c] = self.categoryDialog.categoryName[c].text()
			
			self.categoryDialog.subList1[c] = self.categoryDialog.categorySub1[c].text()
			self.categoryDialog.categorySub1[c].setText(str(self.categoryDialog.subList1[c]).strip("(',')"))#Dialog sub1 SETTEXT
		
			self.categoryDialog.subList2[c] = self.categoryDialog.categorySub2[c].text()
			self.categoryDialog.categorySub2[c].setText(str(self.categoryDialog.subList2[c]).strip("(',')"))#Dialog sub1 SETTEXT
		
			self.categoryDialog.subList3[c] = self.categoryDialog.categorySub3[c].text()
			self.categoryDialog.categorySub3[c].setText(str(self.categoryDialog.subList3[c]).strip("(',')"))#Dialog sub1 SETTEXT  
			
			c+=1

		self.categoryDialog.destroy()
					
	def categorySignalClicked(self): # CONNECT CATEGORIES BUTTONS SIGNAL
	
		self.menuProd_addCat.clicked.connect(self.categorieCreatorSignal)#SIGNAL : CATEGORY DIALOG
		
		#SIGNAL : MENU CATEGORY CLICKED
		i = 0
		while i<=9:
			self.menuProd_categoryList[i].clicked.connect(self.categorySlotClicked)
			i+=1
		
	def categorySlotClicked(self):#SLOT : SET TO TABWIDGET : TabText, VALUES TO BUTTONS, FROM DB , TEXT FROM DIALOG to catNames
				
		pd=0
		while pd <= 74 :		
			self.prod_List[pd].setText("")
			pd+=1
		
		c = 0		
		while c <= 9 :
			if self.sender() == self.menuProd_categoryList[c] :
				self.categoryIndicator.setText(str(self.categoryDialog.catList[c]).strip("(',')"))
				
				if str(self.categoryDialog.subList1[c]).strip("(',')") != "None":
					self.tabWidget.setTabText(0,str(self.categoryDialog.subList1[c]).strip("(',')"))#TAB 1 SETTEXT
				else :
					self.tabWidget.setTabText(0,"")
					
				if str(self.categoryDialog.subList2[c]).strip("(',')") != "None":
					self.tabWidget.setTabText(1,str(self.categoryDialog.subList2[c]).strip("(',')"))#TAB 2 SETTEXT
				else :
					self.tabWidget.setTabText(1,"")
					
				if str(self.categoryDialog.subList3[c]).strip("(',')") != "None":
					self.tabWidget.setTabText(2,str(self.categoryDialog.subList3[c]).strip("(',')"))#TAB 3 SETTEXT
				else :
					self.tabWidget.setTabText(2,"")
			
				query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList1[c]).strip("(',')")+"'")
				subListS1 = list()
				subListS1 = query.fetchall()
				query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList2[c]).strip("(',')")+"'")
				subListS2 = list()
				subListS2 = query.fetchall()
				query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList3[c]).strip("(',')")+"'")
				subListS3 = list()
				subListS3 = query.fetchall()
			
				query.execute("SELECT product_price FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList1[c]).strip("(',')")+"'")
				subListP1 = list()
				subListP1 = query.fetchall()
				query.execute("SELECT product_price FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList2[c]).strip("(',')")+"'")
				subListP2 = list()
				subListP2 = query.fetchall()
				query.execute("SELECT product_price FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList3[c]).strip("(',')")+"'")
				subListP3 = list()
				subListP3 = query.fetchall()
				
				p1=0
				p2=25
				p3=50
				for prod in subListS1 :					
					self.prod_List[p1].setText(str(prod).strip("(',')"))
					self.prod_List[p1].setToolTip(str(prod).strip("(',')")+"\n"+str(subListP1[p1]).strip("(',')")+" DA") 
					p1+=1
				p=0
				for prod in subListS2 :		
					self.prod_List[p2].setText(str(prod).strip("(',')"))
					self.prod_List[p2].setToolTip(str(prod).strip("(',')")+"\n"+str(subListP2[p]).strip("(',')")+" DA") 
					p2+=1
					p+=1
				p=0
				for prod in subListS3 :	
					self.prod_List[p3].setText(str(prod).strip("(',')"))
					self.prod_List[p3].setToolTip(str(prod).strip("(',')")+"\n"+str(subListP3[p]).strip("(',')")+" DA") 
					p3+=1
					p+=1
		
			c+=1

#=========================================			
	def firstView(self):#FIRST VIEW (INDEX) print
		
		self.categoryIndicator.setText(str(self.categoryDialog.catList[0]).strip("(',')"))
		if str(self.categoryDialog.subList1[0]).strip("(',')") != "None":
			self.tabWidget.setTabText(0,str(self.categoryDialog.subList1[0]).strip("(',')"))#TAB 1 SETTEXT
		if str(self.categoryDialog.subList2[0]).strip("(',')") != "None":
			self.tabWidget.setTabText(1,str(self.categoryDialog.subList2[0]).strip("(',')"))#TAB 2 SETTEXT
		if str(self.categoryDialog.subList3[0]).strip("(',')") != "None":
			self.tabWidget.setTabText(2,str(self.categoryDialog.subList3[0]).strip("(',')"))#TAB 3 SETTEXT
			
		query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList1[0]).strip("(',')")+"'")
		self.subListS1 = list()
		self.subListS1 = query.fetchall()
		query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList2[0]).strip("(',')")+"'")
		self.subListS2 = list()
		self.subListS2 = query.fetchall()
		query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList3[0]).strip("(',')")+"'")
		self.subListS3 = list()
		self.subListS3 = query.fetchall()
			
		query.execute("SELECT product_price FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList1[0]).strip("(',')")+"'")
		subListP1 = list()
		subListP1 = query.fetchall()
		query.execute("SELECT product_price FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList2[0]).strip("(',')")+"'")
		subListP2 = list()
		subListP2 = query.fetchall()
		query.execute("SELECT product_price FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList3[0]).strip("(',')")+"'")
		subListP3 = list()
		subListP3 = query.fetchall()
		
		p1=0
		p2=25
		p3=50
		# query.execute("SELECT product_name FROM Products")
		for prod in self.subListS1 :		
			self.prod_List[p1].setText(str(prod).strip("(',')"))
			self.prod_List[p1].setToolTip(str(prod).strip("(',')")+"\n"+str(subListP1[p1]).strip("(',')")+" DA")  
			p1+=1
		p=0	
		for prod in self.subListS2 :		
			self.prod_List[p2].setText(str(prod).strip("(',')"))
			self.prod_List[p2].setToolTip(str(prod).strip("(',')")+"\n"+str(subListP2[p]).strip("(',')")+" DA") 
			p2+=1
			p+=1
		p=0	
		for prod in self.subListS3 :		
			self.prod_List[p3].setText(str(prod).strip("(',')"))
			self.prod_List[p3].setToolTip(str(prod).strip("(',')")+"\n"+str(subListP3[p]).strip("(',')")+" DA") 
			p3+=1
			p+=1

#=========================================				
	def productsSignalClicked(self)	: # CONNECT CATEGORIES BUTTONS SIGNAL
		
		i = 0
		while i <= 74:
			self.prod_List[i].clicked.connect(self.productsSlotClicked)
			i+=1
			
		self.editProduct.clicked.connect(self.productEditSlot)
		
	def productsSlotClicked(self) :	#SLOT : SEE WHO, CREATE variable = prodQntList[i], on click prodNum+=1, FILL TABLEWIDGET
		sender = self.sender()
		pc1 = 0
		
		for prod in self.productDialog.prodNameList :#CATEGORY 1
			
			if sender.text() == str(prod).strip("(',')") and str(prod).strip("(',')") != '' :#and self.prodCatList[pc1] == self.catList[0]:#CATEGORY 1
			
				#GETTING QNT
				self.quantite = str(self.productDialog.prodQntList[pc1]).strip("(,)") 
				self.quantite = int(self.quantite) + 1
				self.productDialog.prodQntList[pc1] = self.quantite
				
				#GETTING price
				self.price = str(self.productDialog.prodPriceList[pc1]).strip("(,)") 
				self.price = int(self.price)
				
				#GETTING total
				self.total = self.quantite * self.price
				self.productDialog.prodTotalList[pc1] =  self.total
				
				#GETTING compta
				self.comptable = str(self.productDialog.prodComptList[pc1]).strip("(,)") 
				self.comptable = int(self.comptable)
				self.comptable +=1
				self.productDialog.prodComptList[pc1] = self.comptable

				#GETTING stock
				self.stockable = str(self.productDialog.prodStockList[pc1]).strip("(,)") 
				self.stockable = int(self.stockable) 
				self.stockable -=1
				self.productDialog.prodStockList[pc1] = self.stockable
				
				self.viewProdNum.setText(str(int(self.viewProdNum.text()) + 1))
				
				self.ProdList.addItem("          "+str(prod).strip("(',')") +"  *  1" )
							
				self.calculTotal()

			pc1+=1

	def calculTotal(self) :	
		self.TotalList = []
		iT = 0
		for prod in self.productDialog.prodNameList :
			if self.productDialog.prodTotalList[iT] != 0  :
				self.totali = str(self.productDialog.prodTotalList[iT]).strip("(',')")
				self.totali=int(self.totali)
				self.TotalList.append(self.totali)
			iT+=1	
		self.viewTotal.display(sum(self.TotalList))	
#=========================================				
	def validationButtonsSignalClicked(self):#CONNECT SELF.VALIDATION BUTTONS TO thEIR SLOT
					
		self.totalButton.clicked.connect(self.validationButtonsSlotClicked)
		self.undoButton.clicked.connect(self.validationButtonsSlotClicked)
		self.chefButton.clicked.connect(self.validationButtonsSlotClicked)
		self.validButton.clicked.connect(self.validationButtonsSlotClicked)
		
	def validationButtonsSlotClicked(self):#VALIDATION, UnDO, 
		sender = self.sender()
		iTu = 0
		
		if sender == self.totalButton :#CREAT A LIST TOTAL, APPEND when ITEM != 0, SUM it, if SUM != 0 --> SELF.Ref++ --> data,DIALOGVLD get text andCALCUL, 
			self.calculTotal()
			if sum(self.TotalList) != 0 :
				
				self.validVLD.VLD_Montant_total.setText(str(sum(self.TotalList)))
				self.validVLD.show()
				
				self.validVLD.VLD_OK.clicked.connect(self.rendurecuSlot)
		
		if sender == self.undoButton :
		
			self.valider = 0
			self.ProdList.clear()
			query.execute("SELECT ref FROM Ref WHERE id = 1")
			
			if str(self.ref) != str(query.fetchone()).strip("(',')") : 	
				self.ref = int(self.ref)-1
				self.viewRefNum.setText(str(self.ref))
				
			self.viewProdNum.setText('0')
					
			self.viewTotal.display(0)
			self.validVLD.VLD_Montant_rendu.setText("0")
			self.validVLD.VLD_Montant_recu.setText("0")
			
			iTu = 0
			for prod in self.productDialog.prodNameList :
				qnt = str(self.productDialog.prodQntList[iTu]).strip("(',')")
				qnt = int(qnt)
					
				self.quantite = str(self.productDialog.prodQntList[iTu]).strip("(,)") 
				self.quantite = int(self.quantite)
				
				if self.quantite != 0:
				
					# self.quantite = 0
					# self.prodQntList[iTu] = self.quantite
					
					self.comptable = str(self.productDialog.prodComptList[iTu]).strip("(,)") 	
					self.comptable = int(self.comptable) - self.quantite	
					self.productDialog.prodComptList[iTu] = self.comptable 
							
					self.stockable = str(self.productDialog.prodStockList[iTu]).strip("(,)")  
					self.stockable = int(self.stockable) + self.quantite	
					self.productDialog.prodStockList[iTu] = self.stockable 	
					
					# GETTING QNT
					self.quantite = 0
					self.productDialog.prodQntList[iTu] = self.quantite
					self.productDialog.prodTotalList[iTu] = 0
					
					self.TotalList = self.productDialog.prodTotalList
				iTu+=1
		
			query.execute("SELECT register_recette_total FROM Register WHERE register_date = '"+self.register.Date+"'")
			self.recetteTest = str(query.fetchone()).strip("(',')")
			
			if str(self.recetteTest).strip("(,)")  != str(self.register.recette).strip("(,)")  :
				self.register.recette = int(self.recetteTest)
				self.register.sumDay.setText(str(self.register.recette))	
			
		if sender == self.chefButton and str(self.viewTotal.intValue()) != '0' :
			
			TICKETchef = Document()
			
			p = TICKETchef.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKETchef.add_heading("\t\t\t\t     "+self.society+"", level=1)
			h.bold = True
			h.italic = True
			p = TICKETchef.add_paragraph("\t\t\t\t    "+self.localDateTime)
			
			p = TICKETchef.add_paragraph(" ___________________________________________________________________________________________________ ")
					
			pi = 0
			tab = TICKETchef.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = ''
			heading_cells[1].text = 'Nom Produit'
			heading_cells[2].text = 'Quantité produit'
			heading_cells[3].text = ''
			
			for prod in self.productDialog.prodNameList :
				qnt = str(self.productDialog.prodQntList[pi]).strip("(',')")
				qnt = int(qnt)
				if qnt != 0:
					cells = tab.add_row().cells
					cells[1].text = str(self.productDialog.prodNameList[pi]).strip("(',')")
					cells[2].text = str(self.productDialog.prodQntList[pi]).strip("(',')")
				pi+=1
			
			p = TICKETchef.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKETchef.add_paragraph(" ___________________________________________________________________________________________________ ")

			TICKETchef.save('DOCUMENTS/TICKETS/Chef/TICKETchef-'+str(self.ref)+'.docx' )
			# print("DocSuccess")
			os.startfile('DOCUMENTS\TICKETS\Chef\TICKETchef-'+str(self.ref)+'.docx' , 'print')	
		
		if sender == self.validButton and str(self.viewTotal.intValue()) != '0' :
			self.ProdList.clear()
			self.register.recette += sum(self.TotalList)
			self.register.sumDay.setText("")
			TICKET = Document()
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKET.add_heading("\t\t\t\t     "+self.society+"", level=1)
			h.bold = True
			h.italic = True
			p = TICKET.add_paragraph("				             "+self.localDateTime)
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
					
			pi = 0
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Nom produit'
			heading_cells[1].text = 'PRIX Produit'
			heading_cells[2].text = 'Quantité produit'
			heading_cells[3].text = 'Montant Produit'
			
			for prod in self.productDialog.prodNameList :
				qnt = str(self.productDialog.prodQntList[pi]).strip("(',')")
				qnt = int(qnt)
				if qnt != 0:
					cells = tab.add_row().cells
					cells[0].text = str(self.productDialog.prodNameList[pi]).strip("(',')")
					cells[1].text = str(self.productDialog.prodPriceList[pi]).strip("(',')") + "DA"
					cells[2].text = str(self.productDialog.prodQntList[pi]).strip("(',')")
					cells[3].text = str(self.productDialog.prodTotalList[pi]).strip("(',')") + "DA"
				pi+=1
				
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			try :	
				p = TICKET.add_heading("\t\t\t\tTotal à payer : "+str(sum(self.TotalList))+" DA ", level=2)
				p.bold = True
			except :
				return 0
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
		
			p = TICKET.add_paragraph("\t\t\t\t Reçu : "+self.validVLD.VLD_Montant_recu.text()+" DA, Rendu : "+self.validVLD.VLD_Montant_rendu.text()+" DA ")
			p = TICKET.add_paragraph("\t\t\t\tCaissier : Annonyme")
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("\t\tTicket No : "+str(self.ref)+". --"+self.society+"-- "+self.numero, level=3)
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

			
			TICKET.save('DOCUMENTS/TICKETS/Tickets/TICKET-'+str(self.ref)+"_"+self.localDate+'.docx' )
			os.startfile('DOCUMENTS\TICKETS\Tickets\TICKET-'+str(self.ref)+"_"+self.localDate+'.docx' , 'print')
		
			self.viewProdNum.setText('0')
			
			query.execute("SELECT ref FROM Ref WHERE id = 1")
					
			self.viewTotal.display(0)
			self.validVLD.VLD_Montant_rendu.setText("0")
			self.validVLD.VLD_Montant_recu.setText("0")
			
			iTu = 0
			for prod in self.productDialog.prodNameList :
					
				self.quantite = str(self.productDialog.prodQntList[iTu]).strip("(,)") 
				self.quantite = int(self.quantite)
				
				if self.quantite != 0:
					
					# GETTING QNT
					self.quantite = 0
					self.productDialog.prodQntList[iTu] = self.quantite
					self.productDialog.prodTotalList[iTu] = 0
				
				iTu+=1
			
			self.ref = int(self.ref)+1
			self.viewRefNum.setText(str(self.ref))
			
			query.execute("SELECT ref FROM Ref WHERE id = 1")
			
			if str(self.ref) != str(query.fetchone()).strip("(',')") : 	
			
				query.execute("UPDATE Ref SET ref = "+str(self.ref)+" WHERE id = 1")
			
			
			query.execute("UPDATE Register SET register_recette_total = "+str(self.register.recette)+" WHERE register_date ='"+self.Date+"'")
			query.execute("SELECT register_recette_total FROM Register WHERE register_date ='"+self.Date+"'")
			
			iTu = 0
			
			for prod in self.productDialog.prodNameList :
			
				self.quantite = str(self.productDialog.prodQntList[iTu]).strip("(,)") 
				self.quantite = int(self.quantite)
				
				if self.quantite != 0:
					
					self.quantite = 0
					self.productDialog.prodQntList[iTu] = self.quantite
				
					self.total = 0
					self.productDialog.prodTotalList[iTu] = self.total
					
				iTu +=1
				
			i=0
			for prod in self.productDialog.prodNameList :
				# query.execute("UPDATE Products SET product_qnt = 0, product_total = 0,product_comptable = '{0}',product_stockable = '{1}' WHERE product_name = '"+str(self.productDialog.prodNameList[i]).strip("(',')")+"'")	
				query.execute("UPDATE Products SET product_qnt= 0 WHERE product_name = '"+str(self.productDialog.prodNameList[i]).strip("(',')")+"'")	
				query.execute("UPDATE Products SET product_comptable =  '"+str(self.productDialog.prodComptList[i]).strip("(',')")+"' WHERE product_name = '"+str(self.productDialog.prodNameList[i]).strip("(',')")+"'")
				query.execute("UPDATE Products SET product_stockable =  '"+str(self.productDialog.prodStockList[i]).strip("(',')")+"' WHERE product_name = '"+str(self.productDialog.prodNameList[i]).strip("(',')")+"'")
				query.execute("UPDATE Products SET product_total = 0 WHERE product_name = '"+str(self.productDialog.prodNameList[i]).strip("(',')")+"'")	
				i+=1
			
			conn.commit()

	def rendurecuSlot(self):
		try :
			self.validVLD.VLD_Montant_rendu.setText(str(int(self.validVLD.VLD_Montant_recu.text()) - sum(self.TotalList) ))
		except :
			return 0 

	#EDIT PRODUCTS	
	def comboInit(self)	:
		
		self.productDialog.prodAdd.clicked.connect(self.prodEditOK)
		self.productDialog.prodModif.clicked.connect(self.prodEditOK)
		self.productDialog.prodDel.clicked.connect(self.prodEditOK)
		
		self.productDialog.categoryAddComboBox.clear()
		self.productDialog.categoryAddComboBox.addItem("...")
		self.productDialog.categoryModifComboBox.clear()
		self.productDialog.categoryModifComboBox.addItem("...")
		self.productDialog.categoryDelComboBox.clear()
		self.productDialog.categoryDelComboBox.addItem("...")
			
		i=0
		while i <= 9:
			self.productDialog.categoryAddComboBox.addItem(self.categoryDialog.categoryName[i].text())
			self.productDialog.categoryModifComboBox.addItem(self.categoryDialog.categoryName[i].text())
			self.productDialog.categoryDelComboBox.addItem(self.categoryDialog.categoryName[i].text())
			i+=1
		
	def productAdder(self):
		try :
			if  self.productDialog.newProdName.text() != '' and self.productDialog.newProdPrice.text() != ''\
			and self.productDialog.categoryAddComboBox.currentText() != ''	and self.productDialog.subCategoryAddComboBox.currentText() != ''\
			and self.productDialog.subCategoryAddComboBox.currentText() != '...' :

				query.execute("INSERT INTO Products (product_category_name,product_subCategory,product_name,product_price,product_qnt,\
				product_comptable,product_stockable,product_total,product_comptable_cst) VALUES\
				('{0}','{1}','{2}','{3}','{4}','{5}','{6}','{7}','{8}')"\
				.format(self.productDialog.categoryAddComboBox.currentText(), self.productDialog.subCategoryAddComboBox.currentText(), \
				self.productDialog.newProdName.text(),self.productDialog.newProdPrice.text(),0,0,0,0,0))
						
				conn.commit()
				self.comboInit()

				self.productDialog.newProdName.setText("")	
				self.productDialog.newProdPrice.setText("")
				self.productDialog.categoryAddComboBox.clear()
				self.productDialog.subCategoryAddComboBox.clear()
				
				self.productDialog.prodDataList()
				 
				self.messageFactory.raiseAdder("Produit")
		
				self.productDialog.categoryAddComboBox.clear()
				i=0
				self.productDialog.categoryAddComboBox.addItem("")
				for prod in self.productDialog.categoryDialog.catList :
					self.productDialog.categoryAddComboBox.addItem(str(prod).strip("(',')"))
					i+=1
					
			else :  
				self.messageFactory.raiseCaseExcept("Toutes les cases")
		except:
			self.messageFactory.raiseCharExcept()

	def productModifier(self):
		try:
			if  self.productDialog.newNameModif.text() != '' and self.productDialog.subCategoryModifComboBox.currentText() != '' and self.productDialog.subCategoryModifComboBox.currentText() != '...' :
					
				if self.productDialog.newNameModif.text() != self.productDialog.prodModifComboList.currentText() :
					query.execute("UPDATE Products SET product_name = '"+str(self.productDialog.newNameModif.text())+"' WHERE product_name = '"+self.productDialog.prodModifComboList.currentText()+"' AND product_category_name = '"+self.productDialog.categoryModifComboBox.currentText()+"'")

				elif self.productDialog.newPriceModif.text() != "" :
					query.execute("UPDATE Products SET product_price = '"+str(self.productDialog.newPriceModif.text())+"' WHERE product_name = '"+self.productDialog.prodModifComboList.currentText()+"' AND product_category_name = '"+self.productDialog.categoryModifComboBox.currentText()+"'")
				
				elif self.productDialog.newNameModif.text() != self.productDialog.prodModifComboList.currentText() and  self.productDialog.newPriceModif.text() != "" :
					query.execute("UPDATE Products SET product_name = '"+str(self.productDialog.newNameModif.text())+"',product_price ='"+self.productDialog.newPriceModif.text()+"' WHERE product_name = '"+self.productDialog.prodModifComboList.currentText()+"' AND product_category_name = '"+self.productDialog.categoryModifComboBox.currentText()+"'")
			
				conn.commit()
				self.comboInit()
				
				self.productDialog.newNameModif.setText("")	
				self.productDialog.newPriceModif.setText("")
				
				self.productDialog.categoryModifComboBox.clear()
				self.productDialog.subCategoryModifComboBox.clear()
				
				self.productDialog.prodDataList()
				
				self.messageFactory.raiseModifier("Produit")
				i=0
				self.productDialog.categoryModifComboBox.addItem("")
				for prod in self.productDialog.categoryDialog.catList :
					self.productDialog.categoryModifComboBox.addItem(str(prod).strip("(',')"))
					i+=1		
			else :  
				self.messageFactory.raiseCaseExcept("toutes les cases")

		except:
			self.messageFactory.raiseCharExcept()

	def productDeleter(self):
		try:
			if  self.productDialog.delProdName.text() != '' and self.productDialog.subCategoryDelComboBox.currentText() != '' and self.productDialog.subCategoryDelComboBox.currentText() != '...' :
																					
				query.execute("DELETE FROM 'Products' WHERE product_name = '"+self.productDialog.delProdName.text()+"' AND product_category_name = '"+self.productDialog.categoryDelComboBox.currentText()+"'")
					
				conn.commit()
				self.comboInit()
			
				self.productDialog.delProdName.setText("")	
				
				self.productDialog.categoryDelComboBox.clear()
				self.productDialog.subCategoryDelComboBox.clear()
				
				self.productDialog.prodDataList()
				
				self.messageFactory.raiseDeleter("Produit")
			
				self.productDialog.categoryDelComboBox.clear()
				i=0
				self.productDialog.categoryDelComboBox.addItem("")
				for prod in self.productDialog.categoryDialog.catList :
					self.productDialog.categoryDelComboBox.addItem(str(prod).strip("(',')"))
					i+=1
				
			else :  
				self.messageFactory.raiseCaseExcept("toutes les cases")
	
		except:
			self.messageFactory.raiseCharExcept()

	def productEditSlot(self):
		self.comboInit()
		self.productDialog.show()
		
		self.productDialog.categoryAddComboBox.currentTextChanged.connect(self.prodSubCatSlot)
		self.productDialog.categoryModifComboBox.currentTextChanged.connect(self.prodSubCatSlot)
		self.productDialog.categoryDelComboBox.currentTextChanged.connect(self.prodSubCatSlot)
			
	def prodEditOK(self):	
		#ADD PRODUCT
		if self.sender() == self.productDialog.prodAdd:
			self.productAdder()
		
		#MODIF PRODUCT
		if self.sender() == self.productDialog.prodModif:
			self.productModifier()
	
		#DELETE PRODUCT
		if self.sender() == self.productDialog.prodDel:
			self.productDeleter()

	def subItemer(self,catCombo,catlist,combo,listsub1,listsub2,listsub3):
			i=0
			combo.clear()
			combo.addItem("...")
			
			while i <= 9:
				if catCombo.currentText() == str(catlist[i]).strip("(',')"):
					
					if listsub1[i].text() != '':
						combo.addItem(str(listsub1[i].text()).strip("(',')"))
						
					if listsub2[i].text() != '':
						combo.addItem(str(listsub2[i].text()).strip("(',')"))
						
					if listsub3[i].text() != '':
						combo.addItem(str(listsub3[i].text()).strip("(',')"))
				i+=1
			combo.currentTextChanged.connect(self.prodSubProdSlot)
				
	def prodSubCatSlot(self):
		
		if self.sender() == self.productDialog.categoryAddComboBox :
			self.subItemer(self.productDialog.categoryAddComboBox, self.categoryDialog.catList,self.productDialog.subCategoryAddComboBox,\
			self.categoryDialog.categorySub1,self.categoryDialog.categorySub2,self.categoryDialog.categorySub3)
		
		if self.sender() == self.productDialog.categoryModifComboBox :
			self.subItemer(self.productDialog.categoryModifComboBox, self.categoryDialog.catList,self.productDialog.subCategoryModifComboBox,\
			self.categoryDialog.categorySub1,self.categoryDialog.categorySub2,self.categoryDialog.categorySub3)
	
		if self.sender() == self.productDialog.categoryDelComboBox :
			self.subItemer(self.productDialog.categoryDelComboBox, self.categoryDialog.catList,self.productDialog.subCategoryDelComboBox,\
			self.categoryDialog.categorySub1,self.categoryDialog.categorySub2,self.categoryDialog.categorySub3)				

	def comboListItemer(self,combolist,subcombo):#JE SUIS ICI	
		#MODIF
		subcombo.clear()
		subcombo.addItem("...")
		
		query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+combolist.currentText()+"'")

		for prod in query.fetchall() :
			subcombo.addItem(str(prod).strip("(',')"))
			
		subcombo.currentTextChanged.connect(self.prodModifSetting)
		
	def prodSubProdSlot(self):
		if self.sender() == self.productDialog.subCategoryModifComboBox :
			self.comboListItemer(self.productDialog.subCategoryModifComboBox,self.productDialog.prodModifComboList)
		if self.sender() == self.productDialog.subCategoryDelComboBox :
			self.comboListItemer(self.productDialog.subCategoryDelComboBox,self.productDialog.prodDelComboList)
			
	def prodModifSetting(self):
		if self.sender() == self.productDialog.prodModifComboList:
			self.productDialog.newNameModif.setText(self.productDialog.prodModifComboList.currentText())
		if self.sender() == self.productDialog.prodDelComboList:
			self.productDialog.delProdName.setText(self.productDialog.prodDelComboList.currentText())
			
#END OF PRODUCT WIDGET TablesNumDialog

#======================================================				
#======================================================
qtCreatorFile = "DESIGN/WIDGETS/mainWindow.ui" # Enter file here.

Ui_MainWindow, QtBaseClass = uic.loadUiType(qtCreatorFile)

class MainWindow(QMainWindow, Ui_MainWindow):
	def __init__(self):
		QMainWindow.__init__(self)
		Ui_MainWindow.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("Restorent Manager System.")
		
		self.declaredUi()
		self.initUi()
		
		self.Date = time.strftime("%Y-%m-%d")
		
	def declaredUi(self):
		self.home = ProductWidget()
		self.register = RegisterCreatorDialog()
		self.settings = SettingCreatorDialog()
		self.connecter = ConnectDialog()
		
		self.registerInit = InitDialog()
		self.registerInit.INIT_OK.clicked.connect(self.closeWindow)
		self.registerInit.INIT_NO.clicked.connect(self.closeWindow)
				
	def initUi(self):
		self.dockWidget.setWidget(self.home)
		self.setCentralWidget(self.dockWidget)
		
		self.EditAct.triggered.connect(self.toolBarCalls)	
		self.HomeAct.triggered.connect(self.toolBarCalls)
		self.CaisseAct.triggered.connect(self.toolBarCalls)
		self.connecter.connOk.clicked.connect(self.toolBarCalls)
		
	def toolBarCalls(self):
	
		if self.sender() == self.HomeAct :
			self.dockWidget.setWidget(self.home)
	
		if self.sender() == self.CaisseAct :
			self.register.totalRegister()
			self.register.show()	
		if self.sender() == self.EditAct :
			self.settings.selection()
			self.connecter.show()
		
		if self.sender() == self.connecter.connOk :
			self.connecter.login(self.settings)
	
	def closeEvent(self, event):
		event.ignore()
		query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
		self.sominit = query.fetchone()
		self.sominit = str(self.sominit).strip("(',')")
		self.registerInit.sumInit.setText(self.sominit)
		self.registerInit.show()
		
	def closeWindow(self):
	
		if self.sender() == self.registerInit.INIT_OK :
			query.execute("UPDATE Register SET register_sum_init = '"+self.registerInit.sumInit.text()+"' WHERE register_date = '"+self.Date+"'")
			query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
			self.sominit = query.fetchone()
			self.sominit = str(self.sominit).strip("(',')")
			self.registerInit.sumInit.setText(self.sominit)
			conn.commit()
			conn.close()
			self.registerInit.destroy()
			self.destroy()
		
		if self.sender() == self.registerInit.INIT_NO :
			self.registerInit.destroy()

class Opener(QDialog):
	def __init__(self):
		QDialog.__init__(self)
		
		self.declaredUi()
		
	def declaredUi(self):
		self.connecter = ConnectDialog()
		self.window = MainWindow()	
		
		self.connecter.connOk.clicked.connect(self.loginSlot)
		
	def loginSlot(self):
		if self.sender() == self.connecter.connOk :
			self.connecter.login(self.window)
						
#============================================================================================================END OPENER
	
if __name__ == "__main__":
	app = QApplication(sys.argv)
	opener = Opener()
	window = opener.connecter
	
	MAC=""
	MAC = uuid.UUID(int=uuid.getnode())
	MAC=str(MAC)
	
	query.execute("SELECT Key FROM MAC WHERE id=1")
	KEY = query.fetchone()
	tester = str(KEY).strip("(',')")

	if tester == '':
		query.execute("UPDATE MAC SET Key = '%s' WHERE id=1" %MAC)
		conn.commit()
		window.show()
		sys.exit(app.exec_())
	
	elif tester == str(MAC):
		window.show()
		sys.exit(app.exec_())
		
	elif tester != str(MAC):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)

		msg.setText("Vous êtes en présence d'une version du produit obselete")
		msg.setInformativeText("La clé de sécurité est incompatible")
		msg.setWindowTitle("ALERT COPY OBSELETE !")
		msg.exec_()

#============================================================================================================END OF PROGRAMME	
